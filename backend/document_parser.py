#!/usr/bin/env python3
"""
document_parser.py - Multi-format Document Parser for AI-CLI
Supports PDF, Word (DOCX), Markdown (MD), and Plain Text (TXT).
"""

import os
import re
from pathlib import Path
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import markdown
    MARKDOWN_SUPPORT = True
except ImportError:
    MARKDOWN_SUPPORT = False


@dataclass
class ParseResult:
    """Result of parsing a document."""
    success: bool
    content: str
    error: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """Document parser supporting multiple formats."""

    FORMAT_ALIASES = {
        "txt": "text",
        "text": "text",
        "plain": "text",
        "md": "markdown",
        "markdown": "markdown",
        "pdf": "pdf",
        "doc": "docx",
        "docx": "docx",
    }

    def __init__(self):
        self.supported_formats = {
            "text": "Plain Text (.txt)",
            "markdown": "Markdown (.md)",
            "pdf": "PDF Document (.pdf)",
            "docx": "Word Document (.docx)",
        }

    def normalize_type(self, document_type: Optional[str], file_path: Optional[str] = None) -> str:
        """Resolve document type from explicit type or file extension."""
        if document_type:
            doc_type = document_type.strip().lower().lstrip(".")
            if doc_type in self.FORMAT_ALIASES:
                return self.FORMAT_ALIASES[doc_type]

        if file_path:
            suffix = Path(file_path).suffix.lower().lstrip(".")
            if suffix in self.FORMAT_ALIASES:
                return self.FORMAT_ALIASES[suffix]

        return "text"

    def parse(self, file_path: str, document_type: Optional[str] = None) -> ParseResult:
        """
        Parse a document from a file path.

        Args:
            file_path: Path to the document file.
            document_type: Optional document type override.

        Returns:
            ParseResult with extracted text content and metadata.
        """
        full_path = Path(file_path).expanduser().resolve()

        if not full_path.exists():
            return ParseResult(
                success=False,
                content="",
                error=f"File not found: {file_path}",
                metadata={}
            )

        resolved_type = self.normalize_type(document_type, str(full_path))

        try:
            if resolved_type == "text":
                content = self.parse_text(str(full_path))
            elif resolved_type == "markdown":
                content = self.parse_markdown(str(full_path))
            elif resolved_type == "pdf":
                content = self.parse_pdf(str(full_path))
            elif resolved_type == "docx":
                content = self.parse_docx(str(full_path))
            else:
                return ParseResult(
                    success=False,
                    content="",
                    error=f"Unsupported document type: {resolved_type}",
                    metadata={}
                )

            metadata = self.extract_metadata(str(full_path), resolved_type, content)
            return ParseResult(success=True, content=content, metadata=metadata)

        except Exception as e:
            return ParseResult(
                success=False,
                content="",
                error=f"Error parsing {resolved_type} document: {str(e)}",
                metadata={"type": resolved_type}
            )

    def parse_content(self, text_content: str, document_type: str = "text") -> ParseResult:
        """Parse raw in-memory content."""
        resolved_type = self.normalize_type(document_type)
        words = len(text_content.split())
        chars = len(text_content)
        lines = len(text_content.splitlines())

        metadata = {
            "type": resolved_type,
            "char_count": chars,
            "word_count": words,
            "line_count": lines,
        }
        return ParseResult(success=True, content=text_content, metadata=metadata)

    def parse_text(self, file_path: str) -> str:
        """Read plain text files supporting multiple encodings."""
        path = Path(file_path)
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]

        for enc in encodings:
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        with open(path, "rb") as f:
            return f.read().decode("utf-8", errors="replace")

    def parse_markdown(self, file_path: str) -> str:
        """Read and normalize markdown content."""
        return self.parse_text(file_path)

    def parse_pdf(self, file_path: str) -> str:
        """Parse PDF document and extract clean text per page."""
        path = Path(file_path)
        if not PDF_SUPPORT:
            try:
                # Fallback to pdfminer if available directly
                import pdfminer.high_level
                return pdfminer.high_level.extract_text(str(path))
            except ImportError:
                raise ImportError("pdfplumber or pdfminer.six is required for PDF parsing. Install with: pip install pdfplumber")

        content_parts = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    content_parts.append(f"--- Página {i} ---\n{text.strip()}")

        if not content_parts:
            return ""

        return "\n\n".join(content_parts)

    def parse_docx(self, file_path: str) -> str:
        """Parse DOCX document extracting paragraphs and tables."""
        path = Path(file_path)
        if not DOCX_SUPPORT:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        doc = docx.Document(path)
        content_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)

        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_cells))
            if table_rows:
                content_parts.append("\n".join(table_rows))

        return "\n\n".join(content_parts)

    def extract_metadata(self, file_path: str, document_type: str, content: str) -> Dict[str, Any]:
        """Extract rich metadata from document."""
        path = Path(file_path)
        stat = path.stat()
        words = len(content.split())
        chars = len(content)
        lines = len(content.splitlines())

        metadata: Dict[str, Any] = {
            "name": path.name,
            "file_name": path.name,
            "size_bytes": stat.st_size,
            "created": stat.st_ctime,
            "modified": stat.st_mtime,
            "type": document_type,
            "word_count": words,
            "char_count": chars,
            "line_count": lines,
        }

        if document_type == "pdf" and PDF_SUPPORT:
            try:
                with pdfplumber.open(path) as pdf:
                    metadata["num_pages"] = len(pdf.pages)
                    if pdf.metadata:
                        for k in ["Title", "Author", "Subject", "Creator", "Producer", "CreationDate"]:
                            if k in pdf.metadata and pdf.metadata[k]:
                                metadata[k.lower()] = str(pdf.metadata[k])
            except Exception:
                pass

        elif document_type == "docx" and DOCX_SUPPORT:
            try:
                doc = docx.Document(path)
                props = doc.core_properties
                metadata["num_paragraphs"] = len(doc.paragraphs)
                metadata["num_tables"] = len(doc.tables)
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
            except Exception:
                pass

        elif document_type == "markdown":
            headings = re.findall(r"^(#{1,6})\s+(.+)$", content, flags=re.MULTILINE)
            metadata["heading_count"] = len(headings)
            metadata["headings"] = [{"level": len(h[0]), "title": h[1].strip()} for h in headings[:10]]

        return metadata


# Global singleton instance
parser = DocumentParser()
