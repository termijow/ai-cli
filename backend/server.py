#!/usr/bin/env python3
"""
server.py - AI-CLI FastAPI Backend Server
Integrates local llama.cpp instance with Document Parser, WhatsApp Analyzer,
Document Generation (Word/PDF), and REST APIs for the AI-CLI Web UI.
"""

import io
import os
import sys
import re
import time
import json
import sqlite3
from datetime import datetime
import logging
from datetime import datetime
import xml.sax.saxutils as saxutils
from pathlib import Path
from typing import Optional, List, Dict, Any, Union
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, JSONResponse
from pydantic import BaseModel, Field
import requests
import uvicorn

# Setup paths and environment
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent

sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(SCRIPT_DIR))

from document_parser import DocumentParser, parser
from whatsapp_parser import WhatsAppParser, whatsapp_parser

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s")
logger = logging.getLogger("ai-cli-backend")

# Try to import python-docx
try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Try to import reportlab
try:
    import reportlab
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfgen import canvas
    import xml.sax.saxutils as saxutils
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

DEFAULT_APA_SYSTEM_PROMPT = """Eres un asistente de redacción y edición académica y profesional de alto nivel.
Cuando respondas o generes contenido, debes estructurarlo SIEMPRE en formato Markdown (.md) limpio, estructurado y riguroso según las siguientes directrices:

1. Estructura Jerárquica:
   - Usa `# Título Principal` para el título del documento o tema central.
   - Usa `## 1. Título de Sección` para cada una de las secciones principales.
   - Usa `### 1.1. Subsección` o `### Caso X:` para casos particulares o subsecciones.
   - Usa `#### Nivel 4` para puntos subordinados.

2. Tablas y Matrices Comparativas:
   - Cuando presentes factores, perfiles, comparativas o datos estructurados, organízalos SIEMPRE como tablas Markdown válidas con encabezados y separadores (| Encabezado 1 | Encabezado 2 | y | :--- | :--- |).
   - NUNCA devuelvas columnas desordenadas en texto plano o tabulaciones sueltas.

3. Listas y Separación de Párrafos:
   - Deja SIEMPRE una línea en blanco antes y después de cada encabezado, lista o tabla.
   - Usa viñetas `- **Concepto:** Explicación` para análisis de casos, justificaciones o listas de características.
   - Usa listas numeradas `1. `, `2. ` para secuencias de pasos o acciones.

4. Formato para Normas APA 7:
   - Mantén un tono formal, académico, objetivo y preciso.
   - Estructura preguntas, respuestas y análisis con subtítulos claros y destacados."""


# Configuration from environment / .env
LLAMA_PORT = int(os.environ.get("LLAMA_PORT", os.environ.get("PORT", 1234)))
LLAMA_HOST = os.environ.get("LLAMA_HOST", "127.0.0.1")
LLAMA_URL = os.environ.get("LLAMA_URL", f"http://{LLAMA_HOST}:{LLAMA_PORT}").rstrip("/")
LLAMA_TIMEOUT = int(os.environ.get("LLAMA_TIMEOUT", 600))  # Default: 10 minutes (600s)

# Documents and Uploads directories
DOCUMENTS_DIR = Path("/app/documents") if Path("/app/documents").exists() else (PROJECT_ROOT / "documents")
DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

UPLOADS_DIR = PROJECT_ROOT / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)


# --- Pydantic Models ---

class SaveDocumentRequest(BaseModel):
    id: str = Field(..., description="Document unique identifier")
    title: Optional[str] = Field("Documento.docx", description="Document title/filename")
    content: Optional[str] = Field("", description="Document content in Markdown")
    type: Optional[str] = Field("docx", description="Document type")
    chatHistory: Optional[List[Dict[str, Any]]] = Field(default_factory=list, description="AI chat history for this document")
    chat_history: Optional[List[Dict[str, Any]]] = Field(None, description="Alias for chatHistory")
    updated_at: Optional[str] = Field(None, description="Last updated ISO timestamp")


class ChatRequest(BaseModel):
    message: str = Field(..., description="User prompt or question")
    context: Optional[str] = Field(None, description="Document or selection text context")
    file: Optional[str] = Field(None, description="Optional document context or filename alias")
    system_prompt: Optional[str] = Field(None, description="Optional system instructions")
    max_tokens: Optional[int] = Field(40960, description="Max generation tokens (full context support up to 40k)")
    temperature: Optional[float] = Field(0.7, description="Sampling temperature")


class ObsidianExportRequest(BaseModel):
    vault_path: Optional[str] = Field(None, description="Custom Obsidian vault directory path")
    chat_file: Optional[str] = Field(None, description="Specific chat file to export")


class DocumentOperation(BaseModel):
    operation: Optional[str] = Field("process", description="Operation: summarize, translate, extract, rephrase")
    document_type: Optional[str] = Field("text", description="Document type: pdf, docx, md, text")
    content: Optional[str] = Field(None, description="Document raw text content")
    text: Optional[str] = Field(None, description="Alternative text payload")
    file_path: Optional[str] = Field(None, description="Path to file on disk")
    language: Optional[str] = Field(None, description="Target language for translation")
    targetLanguage: Optional[str] = Field(None, description="Target language alias")
    sourceLanguage: Optional[str] = Field(None, description="Source language alias")
    prompt: Optional[str] = Field(None, description="Custom prompt instructions")
    extract_type: Optional[str] = Field(None, description="Extraction mode: entities, dates, numbers, summary")
    type: Optional[str] = Field(None, description="Extraction type alias")
    length: Optional[str] = Field("medium", description="Summary length: short, medium, long")
    format: Optional[str] = Field("text", description="Summary format: text, markdown")
    max_tokens: Optional[int] = Field(2048, description="Maximum tokens")


class DocxGenerationRequest(BaseModel):
    title: Optional[str] = Field("Documento", description="Document title")
    content: Optional[str] = Field("", description="Document body content")
    prompt: Optional[str] = Field(None, description="Optional prompt to generate content via LLM")
    output_path: Optional[str] = Field(None, description="Optional output file path")
    sections: Optional[List[Dict[str, Any]]] = Field(None, description="Optional structured sections")


class PdfGenerationRequest(BaseModel):
    title: Optional[str] = Field("Documento", description="Document title")
    content: Optional[str] = Field("", description="Document content (markdown supported)")
    prompt: Optional[str] = Field(None, description="Optional prompt to generate content via LLM")
    format: Optional[str] = Field("markdown", description="Format: markdown or text")
    output_path: Optional[str] = Field(None, description="Optional output file path")


class WhatsAppAnalyzeRequest(BaseModel):
    chat_text: str = Field(..., description="Raw text exported from WhatsApp")
    prompt_override: Optional[str] = Field(None, description="Optional custom extraction prompt")


# --- Lifespan Context Manager ---

@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("🚀 AI-CLI Backend starting up...")
    logger.info(f"LLM Server Target: {LLAMA_URL}")
    yield
    logger.info("🛑 AI-CLI Backend shut down successfully.")


# --- FastAPI Application ---

app = FastAPI(
    title="AI-CLI Backend API",
    description="Unified Backend API for AI Document Processing, WhatsApp Analysis & LLM Services",
    version="1.0.0",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# --- LLM Helper Functions ---

def record_savings(prompt_tokens: int, completion_tokens: int, query_type: str = "web_chat"):
    """Record token usage and savings vs Claude Fable 5 to SQLite and ~/.ai_cli_savings."""
    if prompt_tokens <= 0 and completion_tokens <= 0:
        return
    try:
        # Precios Claude Fable 5: $10/1M input ($0.0000100), $50/1M output ($0.0000500)
        input_cost = prompt_tokens * 0.0000100
        output_cost = completion_tokens * 0.0000500
        total_saving = round(input_cost + output_cost, 4)

        savings_file = Path.home() / ".ai_cli_savings"
        current_savings = 0.0
        if savings_file.exists():
            try:
                current_savings = float(savings_file.read_text().strip())
            except Exception:
                current_savings = 0.0
        new_savings = round(current_savings + total_saving, 2)
        savings_file.write_text(f"{new_savings:.2f}\n")

        db_file = Path.home() / ".ai_cli_db.db"
        if db_file.exists():
            with sqlite3.connect(str(db_file)) as conn:
                conn.execute(
                    "INSERT INTO usage_logs (input_tokens, output_tokens, input_cost, output_savings, total_savings) VALUES (?, ?, ?, ?, ?)",
                    (prompt_tokens, completion_tokens, input_cost, output_cost, new_savings)
                )

        history_dir = Path.home() / ".ai_cli_history"
        history_dir.mkdir(parents=True, exist_ok=True)
        history_file = history_dir / "queries.jsonl"
        entry = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "query_type": query_type,
            "prompt_tokens": prompt_tokens,
            "completion_tokens": completion_tokens,
            "savings": total_saving,
            "total_savings": new_savings
        }
        with open(history_file, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry) + "\n")
    except Exception as e:
        logger.warning(f"No se pudo registrar ahorro: {e}")

def query_llama_cpp(
    prompt: str,
    system_prompt: Optional[str] = None,
    max_tokens: int = 2048,
    temperature: float = 0.7
) -> Dict[str, Any]:
    """
    Call local llama.cpp server OpenAI-compatible chat endpoint.
    Fallback to completions endpoint if chat endpoint is not available.
    """
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    else:
        messages.append({"role": "system", "content": "Eres un asistente de IA inteligente, preciso y servicial."})

    messages.append({"role": "user", "content": prompt})

    chat_url = f"{LLAMA_URL}/v1/chat/completions"
    payload = {
        "model": "localmodel",
        "messages": messages,
        "max_tokens": min(max_tokens, 40960),
        "temperature": temperature,
        "stream": False
    }

    try:
        resp = requests.post(chat_url, json=payload, timeout=LLAMA_TIMEOUT)
        if resp.status_code == 200:
            data = resp.json()
            content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
            usage = data.get("usage", {})
            prompt_tokens = usage.get("prompt_tokens", 0)
            completion_tokens = usage.get("completion_tokens", 0)
            tokens_used = prompt_tokens + completion_tokens
            record_savings(prompt_tokens, completion_tokens, query_type="chat")
            return {"content": content, "tokens_used": tokens_used}
    except Exception as e:
        logger.warning(f"Chat completions failed on {chat_url}: {e}, trying completions fallback...")

    # Fallback to /v1/completions
    compl_url = f"{LLAMA_URL}/v1/completions"
    compl_payload = {
        "model": "localmodel",
        "prompt": f"{system_prompt or ''}\n\nUser: {prompt}\n\nAssistant:",
        "max_tokens": min(max_tokens, 40960),
        "temperature": temperature
    }
    try:
        resp = requests.post(compl_url, json=compl_payload, timeout=LLAMA_TIMEOUT)
        if resp.status_code == 200:
            data = resp.json()
            content = data.get("choices", [{}])[0].get("text", "")
            usage = data.get("usage", {})
            prompt_tokens = usage.get("prompt_tokens", 0)
            completion_tokens = usage.get("completion_tokens", 0)
            tokens_used = prompt_tokens + completion_tokens
            record_savings(prompt_tokens, completion_tokens, query_type="chat")
            return {"content": content, "tokens_used": tokens_used}
    except Exception as e:
        logger.error(f"Completions fallback failed on {compl_url}: {e}")

    # Return mock or offline message if LLM is not responding
    return {
        "content": "⚠️ No se pudo comunicar con el servidor de LLM (llama-server). Verifica que esté iniciado en el puerto 1234 con 'ai serve' o 'ai services start'.",
        "tokens_used": 0
    }


def parse_inline_markdown_runs(text: str):
    """
    Parses a string with inline markdown (**bold**, *italic*, ***both***, `code`)
    and returns a list of (token_text, is_bold, is_italic, is_code).
    """
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)|(\*\*(.+?)\*\*)|(\*(.+?)\*)|(`([^`]+)`)'
    )
    runs = []
    last_idx = 0

    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_idx:
            runs.append((text[last_idx:start], False, False, False))

        full, b_i_txt, _, b_txt, _, i_txt, _, c_txt = match.groups()
        if b_i_txt:
            runs.append((b_i_txt, True, True, False))
        elif b_txt:
            runs.append((b_txt, True, False, False))
        elif i_txt:
            runs.append((i_txt, False, True, False))
        elif c_txt:
            runs.append((c_txt, False, False, True))

        last_idx = end

    if last_idx < len(text):
        runs.append((text[last_idx:], False, False, False))

    return runs if runs else [(text, False, False, False)]


def add_formatted_runs_docx(paragraph, text: str, font_name="Times New Roman", font_size=12, default_bold=False, default_italic=False, color_rgb=(0,0,0)):
    """Add runs to a python-docx paragraph with inline markdown formatting support."""
    tokens = parse_inline_markdown_runs(text)
    for token_text, is_bold, is_italic, is_code in tokens:
        run = paragraph.add_run(token_text)
        run.font.name = 'Courier New' if is_code else font_name
        run.font.size = Pt(font_size - 1 if is_code else font_size)
        run.font.bold = default_bold or is_bold
        run.font.italic = default_italic or is_italic
        run.font.color.rgb = RGBColor(*color_rgb)


def set_apa_table_borders_docx(table):
    """Applies strict APA 7 borders to a python-docx table (Top, Header-Bottom, Bottom horizontal lines; no vertical lines)."""
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '8')  # 1 pt
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tblBorders.append(top)
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)
    
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')  # 0.5 pt
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), '333333')
    tblBorders.append(insideH)

    for b in ['left', 'right', 'insideV']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'none')
        tblBorders.append(node)

    tblPr.append(tblBorders)


def add_apa_header_page_number_docx(section):
    """Adds right-aligned page number to DOCX section header according to APA 7."""
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run._r.append(fldSimple)


def create_docx_bytes(title: str, content: str, sections: Optional[List[Dict[str, Any]]] = None, font_name: str = "Times New Roman") -> bytes:
    """
    Create DOCX document strictly formatted according to APA 7th Edition:
    - 1 inch margins (2.54 cm).
    - Times New Roman 12 pt typography with 1.5 line spacing.
    - Page number at top-right of header.
    - APA 7 Headings (H1 Centered Bold, H2 Left Bold, H3 Left Bold-Italic).
    - APA 7 Tables (Top, Header-Bottom, and Bottom borders only; no vertical lines).
    - Full inline markdown (**bold**, *italic*, `code`) support.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")

    doc = docx.Document()

    # Configure 1 inch margins (APA 7)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        add_apa_header_page_number_docx(section)

    # Style: Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = font_name
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)

    lines = content.splitlines() if content else []

    if sections:
        for section_data in sections:
            sec_title = section_data.get("title") or section_data.get("header")
            if sec_title:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                add_formatted_runs_docx(p, sec_title, font_name=font_name, font_size=13, default_bold=True)
            sec_content = section_data.get("content") or section_data.get("text", "")
            if sec_content:
                for s_line in str(sec_content).splitlines():
                    if s_line.strip():
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = 1.5
                        add_formatted_runs_docx(p, s_line.strip(), font_name=font_name, font_size=12)

    is_first_real_line = True
    i = 0

    while i < len(lines):
        raw_line = lines[i].strip()
        if not raw_line:
            i += 1
            continue

        # Ignore dummy "Documento_1" / "Documento_2" line
        if re.match(r"^Documento[_\s]\d+$", raw_line, re.IGNORECASE):
            i += 1
            continue

        # Page break
        if raw_line in ("---", "***", "___", "[Salto de página]", "[Salto de pagina]"):
            doc.add_page_break()
            i += 1
            continue

        # Markdown Table (| Col 1 | Col 2 |) or Tab-separated table
        is_md_table = raw_line.startswith("|") and raw_line.endswith("|") and raw_line.count("|") >= 2
        is_tsv_table = "\t" in raw_line and i + 1 < len(lines) and "\t" in lines[i+1]

        if is_md_table or is_tsv_table:
            table_lines = []
            if is_md_table:
                while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
            else:
                while i < len(lines) and "\t" in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1

            table_rows = []
            for t_line in table_lines:
                if re.match(r"^\|(\s*:?-+:?\s*\|)+$", t_line):
                    continue  # Skip |---|---| separator
                if is_md_table:
                    cells = [c.strip() for c in t_line.strip("|").split("|")]
                else:
                    cells = [c.strip() for c in t_line.split("\t")]
                if any(cells):
                    table_rows.append(cells)

            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                for r in table_rows:
                    while len(r) < num_cols:
                        r.append("")

                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_apa_table_borders_docx(table)

                for r_idx, row_data in enumerate(table_rows):
                    for c_idx, cell_value in enumerate(row_data):
                        cell = table.cell(r_idx, c_idx)
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.15
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        add_formatted_runs_docx(
                            p,
                            cell_value,
                            font_name=font_name,
                            font_size=10,
                            default_bold=(r_idx == 0)
                        )
                p_after = doc.add_paragraph()
                p_after.paragraph_format.space_after = Pt(6)

            is_first_real_line = False
            continue

        # APA H1 (# Title or Document Title) -> Centered, Bold
        if raw_line.startswith("# ") and not raw_line.startswith("## "):
            text = raw_line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_runs_docx(p, text, font_name=font_name, font_size=14, default_bold=True)
            is_first_real_line = False
            i += 1
            continue

        # APA H2 (## Section or "1. Section", "2. Section") -> Left-aligned, Bold
        if (raw_line.startswith("## ") and not raw_line.startswith("### ")) or (re.match(r"^\d+\.\s+[A-ZÁÉÍÓÚÑ]", raw_line) and len(raw_line) < 120 and not raw_line.endswith(".")):
            text = raw_line[3:].strip() if raw_line.startswith("## ") else raw_line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs_docx(p, text, font_name=font_name, font_size=13, default_bold=True)
            is_first_real_line = False
            i += 1
            continue

        # APA H3 (### Subsection or "1.1. Subsection") -> Left-aligned, Bold, Italic
        if raw_line.startswith("### ") or (re.match(r"^\d+\.\d+\.?\s+[A-ZÁÉÍÓÚÑ]", raw_line) and len(raw_line) < 120 and not raw_line.endswith(".")):
            text = raw_line[4:].strip() if raw_line.startswith("### ") else raw_line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs_docx(p, text, font_name=font_name, font_size=12, default_bold=True, default_italic=True)
            is_first_real_line = False
            i += 1
            continue

        # APA H4 (#### Sub-subsection) -> Left Indent 0.5 in, Bold
        if raw_line.startswith("#### "):
            text = raw_line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs_docx(p, text, font_name=font_name, font_size=12, default_bold=True)
            is_first_real_line = False
            i += 1
            continue

        # Blockquote (> Cita) -> Left Indent 0.5 in
        if raw_line.startswith("> "):
            text = raw_line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_runs_docx(p, text, font_name=font_name, font_size=11, default_italic=True)
            is_first_real_line = False
            i += 1
            continue

        # Bullet list item (- Item or * Item)
        if re.match(r"^[-*•]\s+", raw_line):
            text = re.sub(r"^[-*•]\s+", "", raw_line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(3)
            add_formatted_runs_docx(p, text, font_name=font_name, font_size=12)
            is_first_real_line = False
            i += 1
            continue

        # Numbered list item (1. Item, 2. Item)
        m_num = re.match(r"^(\d+)\.\s+(.*)$", raw_line)
        if m_num:
            num_str, text = m_num.groups()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(f"{num_str}. ").bold = True
            add_formatted_runs_docx(p, text, font_name=font_name, font_size=12)
            is_first_real_line = False
            i += 1
            continue

        # Main title fallback if first line
        if is_first_real_line and len(raw_line) < 120 and not raw_line.endswith('.'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
            add_formatted_runs_docx(p, raw_line, font_name=font_name, font_size=14, default_bold=True)
            is_first_real_line = False
            i += 1
            continue

        # Standard Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        add_formatted_runs_docx(p, raw_line, font_name=font_name, font_size=12)
        is_first_real_line = False
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --- ReportLab APA 7 PDF Implementation ---

if REPORTLAB_AVAILABLE:
    class NumberedCanvasAPA(canvas.Canvas):
        """Draws APA 7 compliant page number at the top-right header."""
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_number(num_pages)
                canvas.Canvas.showPage(self)
            canvas.Canvas.save(self)

        def draw_page_number(self, page_count):
            self.saveState()
            self.setFont("Times-Roman", 10)
            self.setFillColor(colors.HexColor("#333333"))
            self.drawRightString(8.5 * 72 - 72, 11 * 72 - 45, f"{self._pageNumber}")
            self.restoreState()


def convert_markdown_inline_to_reportlab(text: str) -> str:
    """Converts inline markdown (**bold**, *italic*, `code`) into safe ReportLab XML tags."""
    escaped = saxutils.escape(text)
    escaped = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', escaped)
    escaped = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', escaped)
    escaped = re.sub(r'\*(.+?)\*', r'<i>\1</i>', escaped)
    escaped = re.sub(r'`([^`]+)`', r'<font face="Courier">\1</font>', escaped)
    return escaped


def create_pdf_bytes(title: str, content: str) -> bytes:
    """
    Create a PDF document strictly formatted according to APA 7th Edition using ReportLab:
    - 1 inch margins (72 pt).
    - Times-Roman typography with 1.5 line height.
    - Running header page number at top-right.
    - APA 7 Headings & Tables (no vertical borders).
    """
    if not REPORTLAB_AVAILABLE:
        raise ImportError("ReportLab no está instalado. Instala con: pip install reportlab")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'APATitle',
        parent=styles['Heading1'],
        fontName='Times-Bold',
        fontSize=15,
        leading=19,
        alignment=TA_CENTER,
        spaceBefore=12,
        spaceAfter=10,
        textColor=colors.black
    )

    h2_style = ParagraphStyle(
        'APAH2',
        parent=styles['Heading2'],
        fontName='Times-Bold',
        fontSize=13,
        leading=17,
        alignment=TA_LEFT,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.black,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'APAH3',
        parent=styles['Heading3'],
        fontName='Times-BoldItalic',
        fontSize=12,
        leading=16,
        alignment=TA_LEFT,
        spaceBefore=10,
        spaceAfter=4,
        textColor=colors.black,
        keepWithNext=True
    )

    h4_style = ParagraphStyle(
        'APAH4',
        parent=styles['Heading4'],
        fontName='Times-Bold',
        fontSize=12,
        leading=16,
        alignment=TA_LEFT,
        leftIndent=36,
        spaceBefore=8,
        spaceAfter=4,
        textColor=colors.black,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'APABody',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11.5,
        leading=17,
        alignment=TA_LEFT,
        spaceAfter=6,
        textColor=colors.black
    )

    bullet_style = ParagraphStyle(
        'APABullet',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11.5,
        leading=16,
        leftIndent=24,
        firstLineIndent=-12,
        spaceAfter=4,
        textColor=colors.black
    )

    quote_style = ParagraphStyle(
        'APAQuote',
        parent=styles['Normal'],
        fontName='Times-Italic',
        fontSize=11,
        leading=15,
        leftIndent=36,
        rightIndent=18,
        spaceAfter=6,
        textColor=colors.HexColor("#222222")
    )

    table_cell_style = ParagraphStyle(
        'APATableCell',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=9.5,
        leading=13,
        textColor=colors.black
    )

    table_header_style = ParagraphStyle(
        'APATableHeader',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=10,
        leading=14,
        textColor=colors.black
    )

    story = []
    lines = content.splitlines() if content else []
    is_first_real_line = True
    i = 0
    printable_width = 8.5 * 72 - 144  # 468 pt

    while i < len(lines):
        raw_line = lines[i].strip()
        if not raw_line:
            i += 1
            continue

        if re.match(r"^Documento[_\s]\d+$", raw_line, re.IGNORECASE):
            i += 1
            continue

        # Page break
        if raw_line in ("---", "***", "___", "[Salto de página]", "[Salto de pagina]"):
            story.append(PageBreak())
            i += 1
            continue

        # Markdown Table or Tab Table
        is_md_table = raw_line.startswith("|") and raw_line.endswith("|") and raw_line.count("|") >= 2
        is_tsv_table = "\t" in raw_line and i + 1 < len(lines) and "\t" in lines[i+1]

        if is_md_table or is_tsv_table:
            table_lines = []
            if is_md_table:
                while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
            else:
                while i < len(lines) and "\t" in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1

            table_rows = []
            for t_line in table_lines:
                if re.match(r"^\|(\s*:?-+:?\s*\|)+$", t_line):
                    continue
                if is_md_table:
                    cells = [c.strip() for c in t_line.strip("|").split("|")]
                else:
                    cells = [c.strip() for c in t_line.split("\t")]
                if any(cells):
                    table_rows.append(cells)

            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                col_width = printable_width / num_cols

                flowable_data = []
                for r_idx, row in enumerate(table_rows):
                    row_cells = []
                    for c_idx in range(num_cols):
                        val = row[c_idx] if c_idx < len(row) else ""
                        formatted_val = convert_markdown_inline_to_reportlab(val)
                        style_to_use = table_header_style if r_idx == 0 else table_cell_style
                        row_cells.append(Paragraph(formatted_val, style_to_use))
                    flowable_data.append(row_cells)

                t = Table(flowable_data, colWidths=[col_width] * num_cols)
                t.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('TOPPADDING', (0,0), (-1,-1), 5),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                    ('LEFTPADDING', (0,0), (-1,-1), 6),
                    ('RIGHTPADDING', (0,0), (-1,-1), 6),
                    ('LINEABOVE', (0,0), (-1,0), 1.0, colors.black),
                    ('LINEBELOW', (0,0), (-1,0), 0.75, colors.black),
                    ('LINEBELOW', (0,-1), (-1,-1), 1.0, colors.black),
                ]))
                story.append(Spacer(1, 6))
                story.append(t)
                story.append(Spacer(1, 8))

            is_first_real_line = False
            continue

        # APA H1
        if raw_line.startswith("# ") and not raw_line.startswith("## "):
            text = raw_line[2:].strip()
            story.append(Paragraph(convert_markdown_inline_to_reportlab(text), title_style))
            is_first_real_line = False
            i += 1
            continue

        # APA H2
        if (raw_line.startswith("## ") and not raw_line.startswith("### ")) or (re.match(r"^\d+\.\s+[A-ZÁÉÍÓÚÑ]", raw_line) and len(raw_line) < 120 and not raw_line.endswith(".")):
            text = raw_line[3:].strip() if raw_line.startswith("## ") else raw_line
            story.append(Paragraph(convert_markdown_inline_to_reportlab(text), h2_style))
            is_first_real_line = False
            i += 1
            continue

        # APA H3
        if raw_line.startswith("### ") or (re.match(r"^\d+\.\d+\.?\s+[A-ZÁÉÍÓÚÑ]", raw_line) and len(raw_line) < 120 and not raw_line.endswith(".")):
            text = raw_line[4:].strip() if raw_line.startswith("### ") else raw_line
            story.append(Paragraph(convert_markdown_inline_to_reportlab(text), h3_style))
            is_first_real_line = False
            i += 1
            continue

        # APA H4
        if raw_line.startswith("#### "):
            text = raw_line[5:].strip()
            story.append(Paragraph(convert_markdown_inline_to_reportlab(text), h4_style))
            is_first_real_line = False
            i += 1
            continue

        # Blockquote
        if raw_line.startswith("> "):
            text = raw_line[2:].strip()
            story.append(Paragraph(convert_markdown_inline_to_reportlab(text), quote_style))
            is_first_real_line = False
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*•]\s+", raw_line):
            text = re.sub(r"^[-*•]\s+", "", raw_line)
            story.append(Paragraph(f"• {convert_markdown_inline_to_reportlab(text)}", bullet_style))
            is_first_real_line = False
            i += 1
            continue

        # Numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*)$", raw_line)
        if m_num:
            num_str, text = m_num.groups()
            story.append(Paragraph(f"<b>{num_str}.</b> {convert_markdown_inline_to_reportlab(text)}", bullet_style))
            is_first_real_line = False
            i += 1
            continue

        # Title fallback if first line
        if is_first_real_line and len(raw_line) < 120 and not raw_line.endswith('.'):
            story.append(Paragraph(convert_markdown_inline_to_reportlab(raw_line), title_style))
            is_first_real_line = False
            i += 1
            continue

        # Normal Paragraph
        story.append(Paragraph(convert_markdown_inline_to_reportlab(raw_line), body_style))
        is_first_real_line = False
        i += 1

    doc.build(story, canvasmaker=NumberedCanvasAPA)
    buffer.seek(0)
    return buffer.getvalue()



# --- API Routes ---

@app.get("/health", tags=["System"])
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "service": "ai-cli-backend",
        "version": "1.0.0",
        "llama_url": LLAMA_URL
    }


@app.get("/model/info", tags=["Model"])
async def get_model_info():
    """Check connection to llama-server and get model metadata."""
    is_online = False
    model_name = os.environ.get("MODEL_FILE", "Qwen3.5-9B-GGUF")
    try:
        resp = requests.get(f"{LLAMA_URL}/v1/models", timeout=3)
        if resp.status_code == 200:
            is_online = True
            models_data = resp.json().get("data", resp.json().get("models", []))
            if models_data and isinstance(models_data, list):
                model_name = models_data[0].get("id", model_name)
    except Exception:
        is_online = False

    return {
        "status": "online" if is_online else "offline",
        "model": model_name,
        "server_url": LLAMA_URL,
        "port": LLAMA_PORT,
        "features": {
            "docx_generation": DOCX_AVAILABLE,
            "pdf_generation": REPORTLAB_AVAILABLE,
            "apa_standards": True,
            "pdf_parsing": True,
            "whatsapp_analysis": True,
            "multi_format_parser": True
        }
    }


@app.post("/chat", tags=["Chat"])
async def chat_endpoint(request: ChatRequest):
    """Chat with the local LLM with full context up to 40k tokens and auto-compaction."""
    ctx = (request.context or request.file or "").strip()

    # Intelligent compaction if context exceeds 160,000 characters (~40k tokens)
    if len(ctx) > 160000:
        logger.info(f"Context size ({len(ctx)} chars, ~{len(ctx)//4} tokens) exceeds 40k threshold. Applying auto-compaction.")
        ctx = ctx[:80000] + "\n\n[... CONTENIDO INTERMEDIO COMPACTADO AUTOMÁTICAMENTE (>40K TOKENS) ...]\n\n" + ctx[-70000:]

    if ctx:
        full_prompt = f"=== CONTEXTO DEL DOCUMENTO O SELECCIÓN ===\n{ctx}\n\n=== INSTRUCCIÓN DEL USUARIO ===\n{request.message}"
    else:
        full_prompt = request.message

    max_tokens = request.max_tokens or 40960

    llm_res = query_llama_cpp(
        prompt=full_prompt,
        system_prompt=request.system_prompt or DEFAULT_APA_SYSTEM_PROMPT,
        max_tokens=max_tokens,
        temperature=request.temperature or 0.7
    )

    return {
        "content": llm_res["content"],
        "response": llm_res["content"],
        "tokens_used": llm_res["tokens_used"]
    }


@app.post("/documents/parse", tags=["Documents"])
async def parse_document_endpoint(
    file: Optional[UploadFile] = File(None),
    file_path: Optional[str] = Form(None),
    content: Optional[str] = Form(None),
    document_type: Optional[str] = Form(None)
):
    """
    Parse an uploaded file or existing text into clean string content with metadata.
    """
    if file:
        file_ext = Path(file.filename).suffix.lower()
        temp_dest = UPLOADS_DIR / file.filename
        file_bytes = await file.read()
        with open(temp_dest, "wb") as f:
            f.write(file_bytes)

        result = parser.parse(str(temp_dest), document_type=file_ext)
        return {
            "success": result.success,
            "filename": file.filename,
            "content": result.content,
            "metadata": result.metadata,
            "error": result.error
        }

    if file_path:
        result = parser.parse(file_path, document_type=document_type)
        return {
            "success": result.success,
            "filename": Path(file_path).name,
            "content": result.content,
            "metadata": result.metadata,
            "error": result.error
        }

    if content:
        result = parser.parse_content(content, document_type=document_type or "text")
        return {
            "success": True,
            "filename": "document.txt",
            "content": result.content,
            "metadata": result.metadata,
            "error": None
        }

    raise HTTPException(status_code=400, detail="Debes proporcionar un archivo, una ruta (file_path) o contenido (content).")


@app.post("/documents/{document_type}/summarize", tags=["Documents"])
async def summarize_document(document_type: str, request: DocumentOperation):
    """Summarize document content or file."""
    text_content = request.content or request.text or ""
    if not text_content and request.file_path:
        parse_res = parser.parse(request.file_path, document_type=document_type)
        text_content = parse_res.content

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="No se proporcionó contenido para resumir.")

    len_desc = {
        "short": "un resumen breve y conciso en 2-3 oraciones o puntos clave",
        "medium": "un resumen equilibrado destacando los puntos principales y conclusiones",
        "long": "un resumen detallado y exhaustivo que cubra todas las secciones importantes"
    }.get(request.length, "un resumen claro y bien estructurado")

    format_desc = "en formato Markdown con viñetas y títulos claros" if request.format == "markdown" else "en texto fluido"

    prompt = f"""Genera {len_desc} del siguiente documento {format_desc}.
{request.prompt or ''}

DOCUMENTO:
{text_content}"""

    llm_res = query_llama_cpp(prompt, max_tokens=request.max_tokens or 2048)
    return {
        "status": "success",
        "summary": llm_res["content"],
        "result": llm_res["content"],
        "tokens_used": llm_res["tokens_used"]
    }


@app.post("/documents/{document_type}/translate", tags=["Documents"])
async def translate_document(document_type: str, request: DocumentOperation):
    """Translate text content to target language."""
    text_content = request.content or request.text or ""
    if not text_content and request.file_path:
        parse_res = parser.parse(request.file_path, document_type=document_type)
        text_content = parse_res.content

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="No se proporcionó texto para traducir.")

    target_lang = request.language or request.targetLanguage or "español"
    prompt = f"""Traduce de forma fluida, natural y precisa el siguiente texto al idioma: {target_lang}.
Mantén el formato original del documento.

TEXTO ORIGINAL:
{text_content}

TRADUCCIÓN:"""

    llm_res = query_llama_cpp(prompt, max_tokens=request.max_tokens or 2048)
    return {
        "status": "success",
        "translatedText": llm_res["content"],
        "result": llm_res["content"],
        "target_language": target_lang,
        "tokens_used": llm_res["tokens_used"]
    }


@app.post("/documents/{document_type}/extract", tags=["Documents"])
async def extract_information(document_type: str, request: DocumentOperation):
    """Extract key facts, entities, dates, or custom structured data from document."""
    text_content = request.content or request.text or ""
    if not text_content and request.file_path:
        parse_res = parser.parse(request.file_path, document_type=document_type)
        text_content = parse_res.content

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="No se proporcionó contenido para extraer información.")

    ext_type = request.extract_type or request.type or "entities"
    instructions = {
        "entities": "Extrae todas las entidades clave (Personas, Organizaciones, Lugares, Productos, Tecnologías) en formato JSON.",
        "dates": "Extrae todas las fechas, plazos, cronogramas y momentos temporales mencionados en formato JSON con su contexto.",
        "numbers": "Extrae todas las métricas, cantidades, precios, porcentajes y cifras numéricas en formato JSON.",
        "summary": "Extrae los 5 puntos clave e insights fundamentales del documento."
    }.get(ext_type, request.prompt or "Extrae la información más relevante en formato JSON estructurado.")

    prompt = f"""{instructions}

DOCUMENTO:
{text_content}

Responde con la información organizada y clara."""

    llm_res = query_llama_cpp(prompt, max_tokens=request.max_tokens or 2048)
    return {
        "status": "success",
        "extractedData": llm_res["content"],
        "result": llm_res["content"],
        "tokens_used": llm_res["tokens_used"]
    }


@app.post("/documents/word/generate", tags=["Document Generation"])
async def generate_word_document_api(request: DocxGenerationRequest):
    """Generate Word (.docx) document using LLM or structured sections."""
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=400, detail="python-docx no está disponible. Instálalo con: pip install python-docx")

    content = request.content or ""
    tokens_used = 0

    if request.prompt and not content:
        llm_res = query_llama_cpp(
            prompt=f"Escribe el contenido completo y bien estructurado para un documento titulado '{request.title}'.\nInstrucciones: {request.prompt}",
            max_tokens=3000
        )
        content = llm_res["content"]
        tokens_used = llm_res["tokens_used"]

    try:
        doc_bytes = create_docx_bytes(
            title=request.title or "Documento Generado",
            content=content,
            sections=request.sections
        )

        filename = f"{Path(request.title or 'documento').stem}.docx"

        if request.output_path:
            out_file = Path(request.output_path)
            out_file.parent.mkdir(parents=True, exist_ok=True)
            with open(out_file, "wb") as f:
                f.write(doc_bytes)

        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        raise HTTPException(status_code=500, detail=f"Error al generar documento Word: {str(e)}")


@app.post("/documents/pdf/generate", tags=["Document Generation"])
async def generate_pdf_document_api(request: PdfGenerationRequest):
    """Generate APA 7th Edition PDF document."""
    content = request.content or ""
    tokens_used = 0

    if request.prompt and not content:
        llm_res = query_llama_cpp(
            prompt=f"Escribe un documento profesional estructurado en formato Markdown (.md) titulado '{request.title}'.\nInstrucciones: {request.prompt}",
            system_prompt=DEFAULT_APA_SYSTEM_PROMPT,
            max_tokens=3000
        )
        content = llm_res["content"]
        tokens_used = llm_res["tokens_used"]

    if request.format == "json":
        full_md = f"# {request.title or 'Documento'}\n\n{content}\n\n---\n*Generado localmente con AI-CLI*"
        return {
            "status": "success",
            "title": request.title,
            "content": full_md,
            "tokens_used": tokens_used
        }

    if not REPORTLAB_AVAILABLE:
        raise HTTPException(status_code=400, detail="ReportLab no está disponible. Instálalo con: pip install reportlab")

    try:
        pdf_bytes = create_pdf_bytes(
            title=request.title or "Documento Generado",
            content=content
        )
        filename = f"{Path(request.title or 'documento').stem}.pdf"

        if request.output_path:
            out_file = Path(request.output_path)
            out_file.parent.mkdir(parents=True, exist_ok=True)
            with open(out_file, "wb") as f:
                f.write(pdf_bytes)

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error al generar documento PDF: {str(e)}")



@app.post("/whatsapp/parse", tags=["WhatsApp"])
async def parse_whatsapp_chat(request: WhatsAppAnalyzeRequest):
    """Parse WhatsApp exported chat transcript into statistics and structured message log."""
    if not request.chat_text.strip():
        raise HTTPException(status_code=400, detail="El texto del chat de WhatsApp está vacío.")

    stats = whatsapp_parser.parse_text(request.chat_text)
    return {
        "status": "success",
        "stats": stats
    }


@app.post("/whatsapp/analyze", tags=["WhatsApp"])
async def analyze_whatsapp_chat(request: WhatsAppAnalyzeRequest):
    """
    Perform deep relationship analysis and entity extraction on WhatsApp chat using local LLM.
    """
    if not request.chat_text.strip():
        raise HTTPException(status_code=400, detail="El texto del chat de WhatsApp está vacío.")

    stats = whatsapp_parser.parse_text(request.chat_text)
    prompt = request.prompt_override or whatsapp_parser.build_analysis_prompt(stats)

    llm_res = query_llama_cpp(
        prompt=prompt,
        system_prompt="Eres un analista de relaciones y extractor de datos estructurados especializado en transcripciones de mensajería.",
        max_tokens=3000,
        temperature=0.3
    )

    return {
        "status": "success",
        "stats": stats,
        "analysis_raw": llm_res["content"],
        "tokens_used": llm_res["tokens_used"]
    }


@app.post("/whatsapp/analyze-stream", tags=["WhatsApp"])
async def analyze_whatsapp_stream_endpoint(request: WhatsAppAnalyzeRequest):
    """
    Stream real-time batch processing progress and accumulated entity extraction via Server-Sent Events (SSE).
    """
    if not request.chat_text.strip():
        raise HTTPException(status_code=400, detail="El texto del chat de WhatsApp está vacío.")

    from whatsapp_analyzer_engine import LargeChatAnalyzer
    engine = LargeChatAnalyzer(chunk_size_messages=150)

    def query_fn(prompt: str):
        return query_llama_cpp(prompt, max_tokens=1500, temperature=0.2)

    async def event_generator():
        try:
            async for event in engine.analyze_chat_stream(request.chat_text, query_fn):
                payload = json.dumps(event, ensure_ascii=False)
                yield f"data: {payload}\n\n"
        except Exception as e:
            err_payload = json.dumps({"type": "error", "message": str(e)}, ensure_ascii=False)
            yield f"data: {err_payload}\n\n"

    from fastapi.responses import StreamingResponse
    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.get("/whatsapp/dossiers", tags=["WhatsApp"])
async def list_whatsapp_dossiers():
    """List all saved dossiers and profiles in ~/.ai_cli_whatsapp."""
    from whatsapp_analyzer_engine import STORAGE_DIR
    dossiers = []
    if STORAGE_DIR.exists():
        for f in STORAGE_DIR.glob("*_dossier.md"):
            stat = f.stat()
            json_path = f.parent / f"{f.stem.replace('_dossier', '_profile')}.json"
            dossiers.append({
                "id": f.stem,
                "title": f.stem.replace("_dossier", "").replace("_", " ").title(),
                "file_md": f.name,
                "size_bytes": stat.st_size,
                "modified": stat.st_mtime,
                "has_json": json_path.exists()
            })
    return {"status": "success", "dossiers": dossiers}


@app.get("/whatsapp/dossiers/{filename}", tags=["WhatsApp"])
async def get_whatsapp_dossier_content(filename: str):
    """Get the markdown content of a saved dossier."""
    from whatsapp_analyzer_engine import STORAGE_DIR
    safe_name = Path(filename).name
    file_path = STORAGE_DIR / safe_name
    if not file_path.exists():
        file_path = STORAGE_DIR / f"{safe_name}_dossier.md"

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Dossier no encontrado.")

    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    return {"status": "success", "filename": file_path.name, "content": content}


@app.post("/whatsapp/export-obsidian", tags=["WhatsApp", "Obsidian"])
async def export_to_obsidian_vault(request: Optional[ObsidianExportRequest] = None):
    """Export parsed WhatsApp profiles and dossiers into an Obsidian Vault CRM."""
    from obsidian_vault_exporter import ObsidianVaultExporter
    from whatsapp_analyzer_engine import STORAGE_DIR

    custom_vault = Path(request.vault_path) if request and request.vault_path else None
    exporter = ObsidianVaultExporter(vault_path=custom_vault)

    # Load all saved json profiles from STORAGE_DIR
    json_files = list(STORAGE_DIR.glob("*_profile.json"))
    if not json_files:
        raise HTTPException(status_code=400, detail="No hay perfiles de WhatsApp analizados todavía. Analiza una conversación primero.")

    all_contacts = []
    total_events = 0

    for jf in json_files:
        try:
            with open(jf, "r", encoding="utf-8") as f:
                profile_data = json.load(f)
            res = exporter.export_profile(profile_data)
            all_contacts.extend(res.get("contacts_exported", []))
            total_events += res.get("events_exported", 0)
        except Exception as e:
            logger.error(f"Error exporting {jf.name} to Obsidian: {e}")

    return {
        "status": "success",
        "vault_path": str(exporter.vault_path),
        "total_contacts": len(set(all_contacts)),
        "contacts": list(set(all_contacts)),
        "events_count": total_events,
        "message": f"✓ Base de datos sincronizada con Obsidian en {exporter.vault_path}"
    }


@app.get("/whatsapp/obsidian-status", tags=["WhatsApp", "Obsidian"])
async def get_obsidian_status():
    """Get status of the WhatsApp Obsidian Vault."""
    from obsidian_vault_exporter import vault_exporter
    return vault_exporter.get_status()


class WhatsAppSyncRequest(BaseModel):
    limit: Optional[int] = 50
    scrolls: Optional[int] = 8
    target_contact: Optional[str] = None
    auto_analyze: Optional[bool] = False


@app.get("/whatsapp/chats", tags=["WhatsApp"])
async def list_available_chats():
    """List all exported WhatsApp .txt chat files in chats/ directory."""
    from backend.whatsapp_session_exporter import CHATS_EXPORT_DIR
    chats = []
    if CHATS_EXPORT_DIR.exists():
        for f in CHATS_EXPORT_DIR.glob("*.txt"):
            stat = f.stat()
            contact_name = f.stem.replace("Chat_de_WhatsApp_con_", "").replace("Chat de WhatsApp con ", "").replace("_", " ")
            chats.append({
                "filename": f.name,
                "path": str(f),
                "contact": contact_name,
                "size_kb": round(stat.st_size / 1024, 1),
                "modified": stat.st_mtime
            })
    return {"status": "success", "total": len(chats), "chats": sorted(chats, key=lambda x: x["modified"], reverse=True)}


@app.get("/whatsapp/chats/{filename}", tags=["WhatsApp"])
async def get_chat_file_content(filename: str):
    """Returns the text content of an exported WhatsApp chat file."""
    from backend.whatsapp_session_exporter import CHATS_EXPORT_DIR
    safe_name = Path(filename).name
    file_path = CHATS_EXPORT_DIR / safe_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Archivo de chat no encontrado.")
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        content = file_path.read_text(encoding="latin-1")
    return {"status": "success", "filename": file_path.name, "content": content}


@app.get("/whatsapp/sync-status", tags=["WhatsApp"])
async def get_whatsapp_sync_status():
    """Provides complete inventory: exported chats, obsidian CRM contacts, events and pending state."""
    from backend.whatsapp_session_exporter import CHATS_EXPORT_DIR
    from backend.obsidian_vault_exporter import vault_exporter

    exported_chats = []
    obsidian_contacts_map = {}

    # Read Obsidian contacts
    if vault_exporter.contacts_dir.exists():
        for md_file in vault_exporter.contacts_dir.glob("*.md"):
            name = md_file.stem
            memory = vault_exporter.read_contact_memory(name) or {}
            obsidian_contacts_map[name.lower()] = {
                "name": name,
                "file": md_file.name,
                "path": str(md_file),
                "cumpleanos": memory.get("cumpleanos"),
                "ubicacion": memory.get("direccion_ubicacion"),
                "profesion": memory.get("profesion_ocupacion"),
                "modified": md_file.stat().st_mtime
            }

    # Read exported .txt chats
    if CHATS_EXPORT_DIR.exists():
        for f in CHATS_EXPORT_DIR.glob("*.txt"):
            stat = f.stat()
            contact_name = f.stem.replace("Chat_de_WhatsApp_con_", "").replace("Chat de WhatsApp con ", "").replace("_", " ").strip()
            obs_info = obsidian_contacts_map.get(contact_name.lower())
            exported_chats.append({
                "filename": f.name,
                "path": str(f),
                "contact": contact_name,
                "size_kb": round(stat.st_size / 1024, 1),
                "modified": stat.st_mtime,
                "has_obsidian": bool(obs_info),
                "obsidian_info": obs_info
            })

    # Read events
    events = []
    if vault_exporter.events_dir.exists():
        for ef in vault_exporter.events_dir.glob("*.md"):
            try:
                txt = ef.read_text(encoding="utf-8")
                fecha = re.search(r"fecha:\s*\"?([^\n\"]+)\"?", txt)
                hora = re.search(r"hora:\s*\"?([^\n\"]+)\"?", txt)
                lugar = re.search(r"lugar:\s*\"?([^\n\"]+)\"?", txt)
                gcal = re.search(r"google_calendar_url:\s*\"?([^\n\"]+)\"?", txt)
                events.append({
                    "title": ef.stem.replace("_", " "),
                    "fecha": fecha.group(1) if fecha else "",
                    "hora": hora.group(1) if hora else "",
                    "lugar": lugar.group(1) if lugar else "",
                    "gcal_url": gcal.group(1) if gcal else "",
                    "filename": ef.name
                })
            except Exception:
                continue

    return {
        "status": "success",
        "total_exported": len(exported_chats),
        "total_obsidian_contacts": len(obsidian_contacts_map),
        "total_events": len(events),
        "vault_path": str(vault_exporter.vault_path),
        "chats": sorted(exported_chats, key=lambda x: x["modified"], reverse=True),
        "obsidian_contacts": list(obsidian_contacts_map.values()),
        "events": events
    }


@app.get("/whatsapp/obsidian-contact/{contact_name}", tags=["WhatsApp", "Obsidian"])
async def get_obsidian_contact_markdown(contact_name: str):
    """Returns markdown content of the contact card stored in Obsidian CRM."""
    from backend.obsidian_vault_exporter import vault_exporter
    slug = vault_exporter._slugify(contact_name)
    target = vault_exporter.contacts_dir / f"{slug}.md"

    if not target.exists():
        for f in vault_exporter.contacts_dir.glob("*.md"):
            if contact_name.lower() in f.stem.lower() or f.stem.lower() in contact_name.lower():
                target = f
                break

    if not target.exists():
        raise HTTPException(status_code=404, detail=f"Contacto '{contact_name}' no registrado en Obsidian.")

    content = target.read_text(encoding="utf-8")
    return {
        "status": "success",
        "contact": contact_name,
        "filename": target.name,
        "path": str(target),
        "content": content
    }


@app.post("/whatsapp/export-all", tags=["WhatsApp"])
async def export_all_whatsapp_chats(req: Optional[WhatsAppSyncRequest] = None):
    """Automatically exports all chats from WhatsApp Web session in batch."""
    from backend.whatsapp_session_exporter import session_exporter
    limit = req.limit if req and req.limit else 100
    scrolls = req.scrolls if req and req.scrolls else 8

    report = await session_exporter.export_all_chats(limit=limit, max_scrolls_per_chat=scrolls)
    return report


@app.post("/whatsapp/export-all-stream", tags=["WhatsApp"])
async def export_all_whatsapp_chats_stream(req: Optional[WhatsAppSyncRequest] = None):
    """Streams real-time progress of WhatsApp Web batch chat exporter."""
    from backend.whatsapp_session_exporter import session_exporter
    limit = req.limit if req and req.limit else 100
    scrolls = req.scrolls if req and req.scrolls else 8

    event_queue = asyncio.Queue()

    async def progress_cb(data):
        await event_queue.put(data)

    async def run_worker():
        try:
            await event_queue.put({"type": "info", "message": "Iniciando navegador Chromium con perfil persistente..."})
            result = await session_exporter.export_all_chats(
                limit=limit,
                max_scrolls_per_chat=scrolls,
                progress_callback=progress_cb
            )
            await event_queue.put({"type": "done", "result": result, "message": "✓ Exportación masiva completada."})
        except Exception as e:
            logger.error(f"Error en export-all-stream: {e}")
            await event_queue.put({"type": "error", "message": f"Error: {str(e)}"})
        finally:
            await event_queue.put(None)

    asyncio.create_task(run_worker())

    async def event_generator():
        while True:
            item = await event_queue.get()
            if item is None:
                break
            yield f"data: {json.dumps(item)}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.get("/documents/list", tags=["Documents"])
async def list_uploaded_documents():
    """List uploaded and cached documents available for processing."""
    docs = []
    if UPLOADS_DIR.exists():
        for f in UPLOADS_DIR.iterdir():
            if f.is_file() and f.suffix.lower() in [".txt", ".md", ".pdf", ".docx"]:
                stat = f.stat()
                docs.append({
                    "name": f.name,
                    "path": str(f),
                    "size": stat.st_size,
                    "type": f.suffix.lower().lstrip("."),
                    "modified": stat.st_mtime
                })
    return {"status": "success", "documents": docs}


@app.get("/documents/workspace", tags=["Documents", "Persistence"])
@app.get("/api/documents", tags=["Documents", "Persistence"])
async def get_workspace_documents():
    """
    Get all workspace documents and their associated AI chat histories from the persistent documents/ directory.
    """
    documents = []
    if DOCUMENTS_DIR.exists():
        # First load JSON state files (contain content + chatHistory)
        json_files = list(DOCUMENTS_DIR.glob("*.json"))
        loaded_ids = set()

        for jf in json_files:
            try:
                with open(jf, "r", encoding="utf-8") as f:
                    doc_data = json.load(f)
                    if isinstance(doc_data, dict) and "id" in doc_data:
                        documents.append(doc_data)
                        loaded_ids.add(doc_data["id"])
            except Exception as e:
                logger.warning(f"Error reading document state from {jf}: {e}")

        # Also load standalone .md or .txt files that might not have a JSON file yet
        for mf in DOCUMENTS_DIR.iterdir():
            if mf.is_file() and mf.suffix.lower() in [".md", ".txt"]:
                # Check if already loaded by title
                doc_title = mf.name
                if not any(d.get("title") == doc_title or d.get("id") == mf.stem for d in documents):
                    try:
                        content = mf.read_text(encoding="utf-8")
                        documents.append({
                            "id": mf.stem,
                            "title": doc_title,
                            "type": "docx" if doc_title.endswith(".docx") else "md",
                            "content": content,
                            "chatHistory": [],
                            "updated_at": str(mf.stat().st_mtime)
                        })
                    except Exception as e:
                        logger.warning(f"Error reading file {mf}: {e}")

    # Fallback to default initial document if folder is empty
    if not documents:
        default_doc = {
            "id": "1",
            "title": "Documento_1.docx",
            "type": "docx",
            "content": "# Informe Ejecutivo sobre la Implementación de Inteligencia Artificial Local\n\n## 1. Introducción y Definición\n### 1.1. Concepto y Alcance\nLa inteligencia artificial local se refiere al despliegue de modelos de aprendizaje automático y redes neuronales en dispositivos físicos dentro de una organización o infraestructura privada, sin depender de servidores centralizados en la nube.\n\n## 2. Beneficios Estratégicos\n- **Privacidad y Seguridad:** Todos los datos se procesan localmente sin salir de la infraestructura.\n- **Optimización de Costos:** Se eliminan tarifas recurrentes por token o suscripción.\n\n| Factor | Solución Local | Solución Nube |\n| :--- | :--- | :--- |\n| Privacidad | Totalmente aislada | Servidores de terceros |\n| Latencia | Inmediata (ROCm) | Dependiente de conexión |",
            "chatHistory": [],
            "updated_at": "init"
        }
        # Save default doc to disk
        try:
            with open(DOCUMENTS_DIR / "1.json", "w", encoding="utf-8") as f:
                json.dump(default_doc, f, ensure_ascii=False, indent=2)
            with open(DOCUMENTS_DIR / "Documento_1.md", "w", encoding="utf-8") as f:
                f.write(default_doc["content"])
        except Exception as e:
            logger.warning(f"Error saving initial document: {e}")
        documents.append(default_doc)

    return {"status": "success", "documents": documents}


@app.post("/documents/save", tags=["Documents", "Persistence"])
@app.post("/api/documents/save", tags=["Documents", "Persistence"])
async def save_workspace_document(doc: SaveDocumentRequest):
    """
    Save document content and its specific AI chat history to the persistent documents/ directory for git tracking.
    """
    try:
        DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
        chat_hist = doc.chatHistory if doc.chatHistory is not None else (doc.chat_history or [])

        # 1. Save JSON metadata (id, title, content, chatHistory, timestamp)
        doc_dict = {
            "id": doc.id,
            "title": doc.title or f"Documento_{doc.id}.docx",
            "type": doc.type or "docx",
            "content": doc.content,
            "chatHistory": chat_hist,
            "updated_at": doc.updated_at or datetime.now().isoformat()
        }
        json_path = DOCUMENTS_DIR / f"{doc.id}.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(doc_dict, f, ensure_ascii=False, indent=2)

        # 2. Save pure Markdown file for git readability (.md)
        clean_name = re.sub(r'[^\w\s-]', '', Path(doc.title or 'documento').stem).strip().replace(' ', '_')
        if not clean_name:
            clean_name = f"document_{doc.id}"
        md_path = DOCUMENTS_DIR / f"{clean_name}.md"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(doc.content or "")

        return {
            "status": "success",
            "id": doc.id,
            "json_path": str(json_path),
            "md_path": str(md_path),
            "message": f"Documento y chat guardados en {DOCUMENTS_DIR.name}/"
        }
    except Exception as e:
        logger.error(f"Error saving document {doc.id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error al guardar documento: {str(e)}")


@app.delete("/documents/workspace/{doc_id}", tags=["Documents", "Persistence"])
@app.delete("/api/documents/{doc_id}", tags=["Documents", "Persistence"])
async def delete_workspace_document(doc_id: str):
    """Delete a workspace document and all its associated files (.md, .json, .docx, .pdf) from disk."""
    try:
        deleted_files = []
        clean_names_to_delete = {doc_id}

        # 1. Inspect JSON file if present to get original title / clean name
        json_path = DOCUMENTS_DIR / f"{doc_id}.json"
        if json_path.exists():
            try:
                with open(json_path, "r", encoding="utf-8") as f:
                    doc_data = json.load(f)
                    if isinstance(doc_data, dict) and "title" in doc_data:
                        raw_stem = Path(doc_data["title"]).stem
                        c_name = re.sub(r'[^\w\s-]', '', raw_stem).strip().replace(' ', '_')
                        if c_name:
                            clean_names_to_delete.add(c_name)
                        clean_names_to_delete.add(raw_stem)
            except Exception as e:
                logger.warning(f"Error reading {json_path} during delete: {e}")

        # 2. Iterate and delete any file in DOCUMENTS_DIR that matches any stem or filename
        if DOCUMENTS_DIR.exists():
            for f in list(DOCUMENTS_DIR.iterdir()):
                if f.is_file():
                    if f.stem in clean_names_to_delete or f.name in clean_names_to_delete:
                        try:
                            f.unlink()
                            deleted_files.append(f.name)
                            logger.info(f"Deleted file from workspace: {f.name}")
                        except Exception as e:
                            logger.warning(f"Failed to delete {f.name}: {e}")

        return {
            "status": "success",
            "id": doc_id,
            "deleted_files": deleted_files,
            "message": f"Documento y archivos ({', '.join(deleted_files) if deleted_files else doc_id}) eliminados de disco."
        }
    except Exception as e:
        logger.error(f"Error deleting document {doc_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error al eliminar documento: {str(e)}")


# --- Google Calendar Endpoints ---
from backend.google_calendar_sync import calendar_sync

class CalendarAuthRequest(BaseModel):
    code: str
    redirect_uri: str = "urn:ietf:wg:oauth:2.0:oob"

class CreateCalendarEventRequest(BaseModel):
    title: str
    start_time: str
    end_time: Optional[str] = None
    location: Optional[str] = ""
    description: Optional[str] = ""
    contact_name: Optional[str] = None

@app.get("/calendar/status", tags=["Calendar"])
async def get_calendar_status():
    """Returns status of Google Calendar integration."""
    return calendar_sync.get_status()

@app.get("/calendar/auth-url", tags=["Calendar"])
async def get_calendar_auth_url(redirect_uri: str = "urn:ietf:wg:oauth:2.0:oob"):
    """Returns OAuth2 authorization URL for Google Calendar."""
    url = calendar_sync.get_auth_url(redirect_uri)
    if not url:
        raise HTTPException(
            status_code=400,
            detail="credentials.json no encontrado. Coloca tus credenciales de Google Cloud en ~/.ai_cli_google/credentials.json o ~/.ai_cli_whatsapp/google_credentials.json"
        )
    return {"status": "success", "auth_url": url}

@app.post("/calendar/authorize", tags=["Calendar"])
async def complete_calendar_auth(req: CalendarAuthRequest):
    """Exchanges an authorization code for a Google Calendar token."""
    res = calendar_sync.complete_auth(req.code, req.redirect_uri)
    if res.get("status") == "error":
        raise HTTPException(status_code=400, detail=res.get("message"))
    return res

@app.post("/calendar/create-event", tags=["Calendar"])
async def create_calendar_event(req: CreateCalendarEventRequest):
    """Creates an event in Google Calendar and syncs it with Obsidian vault."""
    res = calendar_sync.create_event(
        title=req.title,
        start_dt=req.start_time,
        end_dt=req.end_time,
        location=req.location or "",
        description=req.description or ""
    )

    if req.contact_name:
        try:
            from backend.obsidian_vault_exporter import vault_exporter
            vault_exporter.export_contact_profile({
                "participantes": [{"nombre": req.contact_name}],
                "eventos_y_compromisos": [{
                    "fecha": req.start_time,
                    "descripcion": req.title,
                    "google_calendar_url": res.get("html_link") or res.get("google_calendar_url", "")
                }]
            }, contact_filter=req.contact_name)
        except Exception as e:
            logger.warning(f"No se pudo anexar evento a la nota de Obsidian: {e}")

    return res

@app.get("/calendar/events", tags=["Calendar"])
async def get_calendar_upcoming_events(max_results: int = 10):
    """Returns upcoming Google Calendar events."""
    return {"status": "success", "events": calendar_sync.list_upcoming_events(max_results)}

# --- Headless WhatsApp QR Code Endpoint ---
@app.get("/whatsapp/qr", tags=["WhatsApp"])
async def get_whatsapp_qr_image():
    """Serves the WhatsApp Web QR code image if waiting for authentication on a headless server."""
    qr_path = Path.home() / ".ai_cli_whatsapp" / "qr.png"
    if not qr_path.exists():
        raise HTTPException(
            status_code=404,
            detail="No hay código QR activo en este momento. La sesión de WhatsApp puede estar ya autenticada."
        )
    return FileResponse(qr_path, media_type="image/png")


# --- Server Runner ---

def main():
    port = int(os.environ.get("BACKEND_PORT", 3094))
    host = os.environ.get("BACKEND_HOST", "0.0.0.0")
    logger.info(f"Starting uvicorn server on {host}:{port}...")
    uvicorn.run(app, host=host, port=port, log_level="info")


if __name__ == "__main__":
    main()
